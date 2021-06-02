# -*- coding: utf-8 -*-
"""
Created on Tue Jun  9 14:10:44 2020.

@author: pielsticker
"""
import os
import pickle
import numpy as np
import json
from matplotlib import pyplot as plt

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

#%%
class SpectraPlot:
    """A nx5 array of plots from a given data set."""

    def __init__(self, data, annots):
        """
        Initiate subplots in a nx5 array where n = data.shape[0]/5.

        Parameters
        ----------
        data : array
            A numpy set with two channels: binding energy and intensity.
        annots : list
            List of annotations.

        Returns
        -------
        None.

        """
        self.data = data
        self.annots = annots
        self.no_of_spectra = self.data.shape[0]

        self.no_of_cols = 5
        self.no_of_rows = int(self.no_of_spectra / self.no_of_cols)
        if (self.no_of_spectra % self.no_of_cols) != 0:
            self.no_of_rows += 1

        self.fig, self.axs = plt.subplots(
            nrows=self.no_of_rows, ncols=self.no_of_cols
        )
        plt.subplots_adjust(
            left=0.125,
            bottom=0.5,
            right=4.8,
            top=self.no_of_rows,
            wspace=0.2,
            hspace=0.2,
        )

    def plot(self):
        """
        Populate the plots with the data.

        Returns
        -------
        fig, axs
            Matplotlib objects.
        """
        for i in range(self.no_of_spectra):
            row, col = int(i / self.no_of_cols), i % self.no_of_cols
            x = self.data[i][:, 0]
            y = self.data[i][:, 1]
            annot = self.annots[i]

            try:
                self.axs[row, col].plot(x, y)
                self.axs[row, col].invert_xaxis()
                self.axs[row, col].set_xlim(np.max(x), np.min(x))
                self.axs[row, col].set_xlabel("Binding energy (eV)")
                self.axs[row, col].set_ylabel("Intensity (arb. units)")
                self.axs[row, col].text(
                    0.025,
                    0.4,
                    annot,
                    horizontalalignment="left",
                    verticalalignment="top",
                    transform=self.axs[row, col].transAxes,
                    fontsize=12,
                )

            except IndexError:
                self.axs[row].plot(x, y)
                self.axs[row].invert_xaxis()
                self.axs[row].set_xlim(np.max(x), np.min(x))
                self.axs[row].set_xlabel("Binding energy (eV)")
                self.axs[row].set_ylabel("Intensity (arb. units)")
                self.axs[row].text(
                    0.025,
                    0.4,
                    annot,
                    horizontalalignment="left",
                    verticalalignment="top",
                    transform=self.axs[row, col].transAxes,
                    fontsize=12,
                )

        return self.fig, self.axs


class ClassDistribution:
    """Class for the distribution of classes in a dataset."""

    def __init__(self, task, data_list):
        """
        Calculate the distibutions of the labels.
        
        If the task is "regression", the average distributions
        are calculated. If the task is "classification", calculate how
        many examples of each class are in the different data ses.
        
        Save the distribution in a dict called 'cd'.
        cd: Dictionary of the format {'all data': dict,
                                      'training data': dict,
                                      'validation data': dict,
                                      'test data': dict}.
        Each of the sub-dicts contains the distribution of the labels
        in the data sub-set.

        Parameters
        ----------
        task : str
            If task == 'regression', an average distribution is 
            calculated.
            If task == 'classification', the distribution of the labels
            across the different data sets is calculated.
        data_list : list
            List of numpy arrays containing labels.

        Returns
        -------
        None.

        """
        self.task = task

        self.cd = {
            "all data": {},
            "training data": {},
            "validation data": {},
            "test data": {},
        }

        for i in range(data_list[0].shape[1]):
            self.cd["all data"][str(i)] = 0
            self.cd["training data"][str(i)] = 0
            self.cd["validation data"][str(i)] = 0
            self.cd["test data"][str(i)] = 0

        if self.task == "classification":
            for i, dataset in enumerate(data_list):
                key = list(self.cd.keys())[i]
                for j in range(dataset.shape[0]):
                    argmax_class = np.argmax(dataset[j, :], axis=0)
                    self.cd[key][str(argmax_class)] += 1

        elif self.task == "regression":
            for i, dataset in enumerate(data_list):
                key = list(self.cd.keys())[i]
                average = list(np.mean(dataset, axis=0))
                self.cd[key] = average

    def plot(self, labels):
        """
        Plot the class distribution. Using the labels list as legend.

        Parameters
        ----------
        labels : list
            List of label values for the legend.

        Returns
        -------
        None.

        """
        fig = plt.figure()
        ax = fig.add_axes([0, 0, 1, 1])
        x = np.arange(len(self.cd.keys())) * 1.5
        data = []

        if self.task == "classification":
            plt.title("Class distribution")
            for k, v in self.cd.items():
                data_list = []
                for key, value in v.items():
                    data_list.append(value)
            data.append(data_list)
            data = np.transpose(np.array(data))

        elif self.task == "regression":
            plt.title("Average distribution across the classes")
            # Plot of the average label distribution in the different
            # data sets.
            for k, v in self.cd.items():
                data.append(v)
            data = np.transpose(np.array(data))

        for i in range(data.shape[0]):
            ax.bar(x + i * 0.25, data[i], align="edge", width=0.2)
        plt.legend(labels)
        plt.xticks(ticks=x + 0.5, labels=list(self.cd.keys()))
        plt.show()


class TrainingGraphs:
    """Class for graphs with the result of the training in Keras."""

    def __init__(self, history, fig_dir):
        """
        Take a dictionary containing the results from training.

        Parameters
        ----------
        history : dict
            # A dictionary containing the results from the training of a
            neural network in Keras.
        fig_dir : str
            The name of the directory where the figures shall be saved.

        Returns
        -------
        None.

        """
        self.history = history
        self.fig_dir = fig_dir

    def plot_loss(self, to_file=True):
        """
        Plot the training and validation loss against the epochs.

        Returns
        -------
        None.

        """
        fig, ax = plt.subplots()
        ax.plot(self.history["loss"], linewidth=3)
        ax.plot(self.history["val_loss"], linewidth=3)
        ax.set_title("Loss")
        ax.set_ylabel("Loss")
        ax.set_xlabel("Epoch")
        ax.legend(["Train", "Validation"])
        fig_name = os.path.join(self.fig_dir, "loss.png")
        if to_file:
            fig.savefig(fig_name)
        plt.show()

    def plot_accuracy(self, to_file=True):
        """
        Plot the training and validation accuracy against the epochs.

        Returns
        -------
        None.

        """
        fig, ax = plt.subplots()
        ax.plot(self.history["accuracy"], linewidth=3)
        ax.plot(self.history["val_accuracy"], linewidth=3)
        ax.set_title("Accuracy")
        ax.set_ylabel("Classification Accuracy")
        ax.set_xlabel("Epoch")
        ax.legend(["Train", "Validation"])
        fig_name = os.path.join(self.fig_dir, "accuracy.png")
        if to_file:
            fig.savefig(fig_name)
        plt.show()


class Report:
    """Report on the results of the training in keras."""

    def __init__(self, dir_name=""):
        """
        Initialize a docx document.
        
        Load the data from the hyperparamters file.

        Parameters
        ----------
        dir_name : str, optional
            The name of the directory where the report shall be saved.
            The default is ''.

        Returns
        -------
        None.

        """
        self.document = Document()
        style = self.document.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(10)

        # Get the data
        root_dir = os.getcwd()
        self.model_dir = os.path.join(*[root_dir, "runs", dir_name, "model",])
        self.log_dir = os.path.join(*[root_dir, "runs", dir_name, "logs"])
        self.fig_dir = os.path.join(*[root_dir, "runs", dir_name, "figures"])

        (
            self.name_data,
            self.train_data,
            self.model_summary,
        ) = self.get_hyperparams()
        self.results = self.get_results()
        self.class_dist = self.results["class_distribution"]
        self.filename = os.path.join(self.log_dir, "report.docx")
        self.create_document()

    def create_document(self):
        """
        Add data from the results to the report.

        Returns
        -------
        None.

        """
        self.document.add_heading("Training report", 0)

        # Add the names and basic information.
        self.document.add_heading("Data:", 1)

        name_table = self.document.add_table(
            rows=len(self.name_data.keys()), cols=2
        )
        for key, value in self.name_data.items():
            j = int(list(self.name_data.keys()).index(key))
            name_table.cell(j, 0).text = key + ":"
            name_table.cell(j, 1).text = str(value)

        self.document.add_heading("Distribution:", 1)
        dist_table = self.document.add_table(
            rows=len(self.class_dist.keys()) + 1,
            cols=len(next(iter(self.class_dist.values()))) + 1,
        )

        dist_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, name in enumerate(self.name_data["Labels"]):
            dist_table.cell(0, i + 1).text = name
        for item, param in self.class_dist.items():
            j = int(list(self.class_dist.keys()).index(item)) + 1
            dist_table.cell(j, 0).text = item

            for key, value in enumerate(self.class_dist[item]):
                k = int(key) + 1
                dist_table.cell(j, k).text = str(np.round(value, 3))

        self.document.add_page_break()

        # Add information about the training parameters
        self.document.add_heading("Training parameters:", 1)

        train_table = self.document.add_table(
            rows=len(self.train_data.keys()), cols=2
        )

        for key, value in self.train_data.items():
            j = int(list(self.train_data.keys()).index(key))
            train_table.cell(j, 0).text = key + ":"
            train_table.cell(j, 1).text = str(value)

        # Add the model architecture
        self.document.add_heading("Model architecture", 1)

        par = self.document.add_paragraph(self.model_summary)
        par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.document.add_page_break()

        # Add loss and accuracy values.
        self.document.add_heading("Loss & accuracy", 1)

        loss_file = os.path.join(self.fig_dir, "loss.png")
        self.document.add_picture(loss_file, width=Cm(12))
        last_paragraph = self.document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            acc_file = os.path.join(self.fig_dir, "accuracy.png")
            self.document.add_picture(acc_file, width=Cm(12))
            last_paragraph = self.document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except FileNotFoundError:
            pass

        # Add results on the test data.
        self.document.add_heading("Results", 1)

        result_table = self.document.add_table(rows=2, cols=2)
        result_table.cell(0, 0).text = "Test loss:"
        result_table.cell(0, 1).text = str(
            np.round(self.results["test_loss"], decimals=3)
        )
        try:
            result_table.cell(1, 0).text = "Test accuracy:"
            result_table.cell(1, 1).text = str(
                np.round(self.results["test_accuracy"], decimals=3)
            )

            for row in result_table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except KeyError:
            pass

        self.document.add_page_break()

        # Add predictions on random data.
        self.document.add_heading("Predictions for 5 random examples", 1)

        self.document.add_heading("Training data", 2)

        r = np.random.randint(0, self.results["y_train"].shape[0] - 5)
        pred_train_5 = self.results["pred_train"][r : r + 5, :]
        y_train_5 = self.results["y_train"][r : r + 5, :]

        p = self.document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = None
        run = p.add_run()
        run.text = "Predictions:"
        run.font.underline = True
        _ = self.add_result_table(pred_train_5)

        p = self.document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = None
        run = p.add_run()
        run.text = "Correct labels:"
        run.font.underline = True
        self.add_result_table(y_train_5)

        self.document.add_heading("Test data", 2)
        s = np.random.randint(0, self.results["y_test"].shape[0] - 5)
        pred_test_5 = self.results["pred_test"][s : s + 5, :]
        y_test_5 = self.results["y_test"][s : s + 5, :]

        p = self.document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = None
        run = p.add_run()
        run.text = "Predictions:"
        run.font.underline = True
        self.add_result_table(pred_test_5)

        p = self.document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = None
        run = p.add_run()
        run.text = "Correct labels:"
        run.font.underline = True
        self.add_result_table(y_test_5)

    def add_result_table(self, data_array):
        """
        Store and display the results from training.

        Parameters
        ----------
        data_array : ndarray
            Array with the results from training.

        Returns
        -------
        None.

        """
        new_table = self.document.add_table(
            rows=data_array.shape[0] + 1, cols=data_array.shape[1]
        )
        for row in new_table.rows:
            row.height = Cm(0.5)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        for i, name in enumerate(self.name_data["Labels"]):
            new_table.cell(0, i).text = name
            new_table.cell(0, i).paragraphs[
                0
            ].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if data_array.dtype == "float32":
            a = np.around(data_array, decimals=4)
            row_sums = a.sum(axis=1)
            data_array = a / row_sums[:, np.newaxis]
            data_array = np.around(data_array, decimals=2)

        for i in range(data_array.shape[0]):
            for j in range(data_array.shape[1]):
                new_table.cell(i + 1, j).text = str(data_array[i, j])
                new_table.cell(i + 1, j).paragraphs[
                    0
                ].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def get_hyperparams(self):
        """
        Load the hyperparameters of the training from the JSON file.

        Returns
        -------
        name_data : dict
            Basic information about the experiment.
        class_distribution : dict
            Distribution of the class in the data sets.
        train_data : dict
            Information about the training parameters..
        model_summary : dict
            Summary of the model in str format.

        """
        hyperparam_file_name = os.path.join(
            self.log_dir, "hyperparameters.json"
        )
        with open(hyperparam_file_name, "r") as json_file:
            data_dict = json.load(json_file)

        name_data = {
            "Name": data_dict["exp_name"],
            "Time created": data_dict["time"],
            "No. of classes": data_dict["num_of_classes"],
            "Labels": data_dict["labels"],
            "Total no. of samples": data_dict["no_of_examples"],
            "Train-test-split": data_dict["train_test_split"],
            "Train-val-split": data_dict["train_val_split"],
            "No. of training samples": data_dict["No. of training samples"],
            "No. of validation samples": data_dict[
                "No. of validation samples"
            ],
            "No. of test samples": data_dict["No. of test samples"],
            "Shape of each sample": data_dict["Shape of each sample"],
        }

        train_data = {
            "Optimizer": data_dict["optimizer"],
            "Learning rate": data_dict["learning_rate"],
            "Loss function": data_dict["loss"],
            "Epochs trained": data_dict["epochs_trained"],
            "Batch size": data_dict["batch_size"],
        }

        model_summary = data_dict["model_summary"]

        return name_data, train_data, model_summary

    def get_results(self):
        """
        Load the results from the pickle file.
        
        Results include e.g. the test data and the predictions.

        Returns
        -------
        data : dict
            Dictionary of numpy arrays with the results..

        """
        file_name = os.path.join(self.log_dir, "results.pkl")

        with open(file_name, "rb") as pickle_file:
            data = pickle.load(pickle_file)

        return data

    def write(self):
        """
        Store the document in the logs folder.

        Returns
        -------
        None.

        """
        self.document.save(self.filename)
        print("Report saved!")


#%%
if __name__ == "__main__":
    dir_name = "20210226_16h07m_Fe_4_classes_linear_comb_new_noise"
    rep = Report(dir_name)
    data = rep.get_results()
    summary = rep.model_summary
    rep.write()
